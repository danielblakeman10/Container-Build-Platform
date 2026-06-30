#!/usr/bin/env python3
"""Build the Container Build Platform tutorial.docx from real repo code."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ASSETS = os.path.dirname(os.path.abspath(__file__))
REPO = os.path.dirname(ASSETS)
OUT = os.path.join(REPO, "tutorial.docx")

DARK = RGBColor(0x1e, 0x29, 0x3b)
BLUE = RGBColor(0x25, 0x63, 0xeb)
GREEN = RGBColor(0x16, 0xa3, 0x4a)
ORANGE = RGBColor(0xea, 0x58, 0x0c)
GRAY = RGBColor(0x47, 0x55, 0x69)
CODEBG = "F3F4F6"

doc = Document()

# Base style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def shade(paragraph, color_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color_hex)
    pPr.append(shd)


def code_block(text):
    """Monospace shaded code block, one paragraph with line breaks."""
    p = doc.add_paragraph()
    shade(p, CODEBG)
    pf = p.paragraph_format
    pf.left_indent = Pt(8); pf.right_indent = Pt(8)
    pf.space_before = Pt(6); pf.space_after = Pt(6)
    pf.line_spacing = 1.0
    lines = text.strip("\n").split("\n")
    for i, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
        # force monospace east-asian too
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts"); rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), "Consolas")
        rfonts.set(qn("w:hAnsi"), "Consolas")
        if i < len(lines) - 1:
            run.add_break()
    return p


def body(text, bold=False, italic=False, color=None, size=11, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Pt(18 + level * 18)
    p.paragraph_format.space_after = Pt(3)
    # support inline bold via **
    parts = text.split("**")
    for i, part in enumerate(parts):
        r = p.add_run(part)
        if i % 2 == 1:
            r.bold = True
    return p


def numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    parts = text.split("**")
    for i, part in enumerate(parts):
        r = p.add_run(part)
        if i % 2 == 1:
            r.bold = True
    return p


def h1(text):
    doc.add_page_break()
    p = doc.add_heading(text, level=1)
    for r in p.runs: r.font.color.rgb = BLUE


def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs: r.font.color.rgb = DARK


def h3(text):
    p = doc.add_heading(text, level=3)
    for r in p.runs: r.font.color.rgb = GRAY


def add_image(name, width=6.3, caption=None):
    path = os.path.join(ASSETS, name)
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption)
        r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GRAY


def callout(label, text, color=ORANGE):
    p = doc.add_paragraph()
    shade(p, "FFF7ED")
    p.paragraph_format.left_indent = Pt(8); p.paragraph_format.right_indent = Pt(8)
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"{label}  ")
    r.bold = True; r.font.color.rgb = color
    r2 = p.add_run(text)
    r2.font.size = Pt(10)


def info_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True; run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(9.5)
    return t


# ============================================================
# COVER
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Container Build Platform")
r.bold = True; r.font.size = Pt(30); r.font.color.rgb = BLUE
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("A Step-by-Step, Low-Level Tutorial")
r.font.size = Pt(16); r.font.color.rgb = GRAY
sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("Enterprise Container Factory \u2014 DevSecOps Build Pipeline\n"
                 "Docker \u00b7 GitHub Actions \u00b7 Trivy \u00b7 AWS ECR \u00b7 Cosign \u00b7 Terraform")
r.font.size = Pt(11); r.italic = True; r.font.color.rgb = DARK

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Domain 1 (DevOps) \u2014 Project 3\n"
                 "Author: Daniel Blakeman\n"
                 "Repository: github.com/danielblakeman10/Container-Build-Platform")
r.font.size = Pt(10); r.font.color.rgb = GRAY

doc.add_paragraph()
add_image("diagram_pipeline.png", width=6.5,
          caption="Figure 0 \u2014 The end-to-end pipeline you will build in this tutorial.")

# ============================================================
# TABLE OF CONTENTS (static)
# ============================================================
h2("What This Tutorial Covers")
body("This document walks you through the Container Build Platform from absolute "
     "zero to a working, secure, automated container build-and-publish pipeline. "
     "Every section explains not just WHAT to type, but WHY it works that way at a "
     "low level \u2014 what each tool does under the hood, what the security trade-offs "
     "are, and how the pieces fit together. The code shown is the actual code in your "
     "repository.")
for i, t in enumerate([
    "Part 1 \u2014 Concepts & Mental Model (what a 'build platform' really is)",
    "Part 2 \u2014 Prerequisites & Local Environment Setup",
    "Part 3 \u2014 The Multi-Stage Dockerfile (line by line)",
    "Part 4 \u2014 Building & Testing Locally",
    "Part 5 \u2014 Image Tagging Strategy & the helper scripts",
    "Part 6 \u2014 Vulnerability Scanning with Trivy",
    "Part 7 \u2014 The GitHub Actions CI Pipeline (job by job)",
    "Part 8 \u2014 AWS Setup: IAM, ECR, S3 & Secrets",
    "Part 9 \u2014 Pushing to ECR & the CodeBuild alternative",
    "Part 10 \u2014 Image Signing & SBOMs with Cosign",
    "Part 11 \u2014 Image Lifecycle Management & Cleanup",
    "Part 12 \u2014 Infrastructure as Code (Terraform)",
    "Part 13 \u2014 End-to-End Run-Through & Verification",
    "Appendix A \u2014 Troubleshooting",
    "Appendix B \u2014 Glossary",
]):
    bullet(t)

# ============================================================
# PART 1 - CONCEPTS
# ============================================================
h1("Part 1 \u2014 Concepts & Mental Model")
h2("1.1 What problem does this platform solve?")
body("Imagine ten different teams each writing their own way to build a Docker image, "
     "tag it, scan it for security holes, and push it to a registry. You get ten "
     "inconsistent, often insecure pipelines. A Container Build Platform (sometimes "
     "called an \u201cEnterprise Container Factory\u201d) centralizes that into one reusable, "
     "audited, secure path that every team plugs into.")
body("Your platform takes source code and produces a signed, scanned, versioned "
     "container image sitting safely in a private registry \u2014 automatically, on every "
     "push. Nothing reaches production without passing through the same gates.")

h2("1.2 The core vocabulary")
info_table(
    ["Term", "Plain-English meaning"],
    [
        ["Image", "A frozen, read-only template containing your app + everything it needs to run."],
        ["Container", "A running instance of an image (like a process started from a program)."],
        ["Registry", "A storage server for images. AWS ECR is Amazon's private registry."],
        ["Layer", "A cached filesystem diff. Each Dockerfile instruction creates one."],
        ["Multi-stage", "Using throwaway build stages so the final image stays tiny & clean."],
        ["CI/CD", "Continuous Integration / Delivery \u2014 automation that builds & ships on every push."],
        ["CVE", "A publicly catalogued security vulnerability (Common Vulnerabilities & Exposures)."],
        ["SBOM", "Software Bill of Materials \u2014 a manifest of everything inside your image."],
        ["Signing", "Cryptographically proving an image is authentic and untampered."],
    ])

h2("1.3 The six pipeline stages")
body("Everything in this project maps to six conceptual stages. Keep this picture in "
     "your head for the rest of the tutorial:")
add_image("diagram_pipeline.png", width=6.5,
          caption="Figure 1 \u2014 Build \u2192 Test \u2192 Scan \u2192 Push \u2192 Sign, triggered by a git push.")
info_table(
    ["#", "Stage", "Tool", "Gate"],
    [
        ["1", "Build", "Multi-stage Dockerfile + Buildx", "Required"],
        ["2", "Test", "Unit/integration tests in build", "Required"],
        ["3", "Scan", "Trivy (CVE + misconfig)", "WARN on High, FAIL on Critical"],
        ["4", "Push", "AWS ECR", "Required"],
        ["5", "Sign", "Cosign / Notary + SBOM", "Recommended"],
        ["6", "Cleanup", "ECR lifecycle policy", "Scheduled"],
    ])

# ============================================================
# PART 2 - PREREQUISITES
# ============================================================
h1("Part 2 \u2014 Prerequisites & Local Environment Setup")
h2("2.1 Tools you need installed")
info_table(
    ["Tool", "Version", "Why"],
    [
        ["Docker", "24.0+", "Builds and runs containers. Needs BuildKit (default in 24+)."],
        ["AWS CLI", "v2", "Talks to ECR, S3, IAM from your terminal."],
        ["Git", "any recent", "Version control + the SHA tags your scripts generate."],
        ["Trivy", "latest", "Local vulnerability scanning (optional but recommended)."],
        ["Cosign", "v2.2+", "Image signing (optional, for the signing stage)."],
    ])
callout("[i] NOTE", "On your Windows machine, run these in Git-Bash so the shell scripts "
        "(.sh files) work. Docker Desktop must be running before any docker command.")

h2("2.2 Verify your setup")
body("Run each of these. Every command should print a version, not an error:")
code_block(
    "docker --version          # Docker version 24.x or higher\n"
    "docker buildx version     # buildx is what enables multi-arch builds\n"
    "aws --version             # aws-cli/2.x\n"
    "git --version\n"
    "trivy --version           # optional\n"
    "cosign version            # optional")

h2("2.3 Clone the repository")
code_block(
    "git clone https://github.com/danielblakeman10/Container-Build-Platform.git\n"
    "cd Container-Build-Platform")
body("Take a moment to look at the folder structure \u2014 each folder has a single, clear "
     "responsibility:")
info_table(
    ["Folder", "Contains", "Covered in"],
    [
        ["dockerfiles/", "The multi-stage Dockerfile template + compose", "Part 3"],
        [".github/workflows/", "ci.yml, ecr.yml, signing.yml pipelines", "Parts 7,10,11"],
        [".github/scripts/", "parse-version, image-tagging, cleanup-ecr", "Parts 5,11"],
        [".github/codebuild-spec.yml", "AWS-native build alternative", "Part 9"],
        ["terraform/", "ECR + lifecycle infrastructure as code", "Part 12"],
        ["registry/", "ECR lifecycle policy JSON", "Part 11"],
        ["examples/", "Sample apps to actually build", "Part 4"],
    ])

# ============================================================
# PART 3 - DOCKERFILE
# ============================================================
h1("Part 3 \u2014 The Multi-Stage Dockerfile (Line by Line)")
body("This is the heart of the platform. Open dockerfiles/Dockerfile. We'll dissect "
     "every instruction and explain what Docker does internally for each.")

h2("3.1 Why multi-stage at all?")
body("A naive Dockerfile ships the compiler, source code, and build tools to production. "
     "That bloats the image (hundreds of MB) and hands attackers a toolbox. Multi-stage "
     "builds use temporary stages to compile, then copy ONLY the finished binary into a "
     "tiny final image. Everything else is thrown away.")
add_image("diagram_multistage.png", width=6.4,
          caption="Figure 2 \u2014 Build artifacts are copied forward; build tooling is discarded.")

h2("3.1.1 IMPORTANT \u2014 a Dockerfile is a recipe, not a script you type")
body("Before we go line by line, clear up the single most common point of confusion. "
     "The block of text below is NOT a list of commands you paste into your terminal "
     "one at a time. It is the entire contents of ONE file in your repository:", bold=False)
code_block("C:\\...\\Container-Build-Platform\\dockerfiles\\Dockerfile")
body("Think of it like a cake recipe. You do not perform each step in your own kitchen. "
     "You hand the whole recipe to a baker \u2014 Docker \u2014 and say \u201cbuild this.\u201d Docker then "
     "reads the file top to bottom and runs every instruction itself, inside a "
     "throwaway container. You never run the FROM / RUN / COPY lines by hand.")
body("That single hand-off is ONE command, run ONCE, from the repository root folder:", bold=True)
code_block("# from C:\\...\\Container-Build-Platform  (the repo root)\n"
           "docker build -f dockerfiles/Dockerfile -t myapp:latest .")
info_table(
    ["You write...", "What actually happens"],
    [
        ["The Dockerfile (the recipe)", "You edit it once in dockerfiles/Dockerfile and save it."],
        ["docker build ... (the hand-off)", "You type this ONE line in Git-Bash. Docker does the rest."],
        ["Each FROM / RUN / COPY line", "Docker runs it FOR you, in order, inside the build. Never you."],
    ])
callout("[i] KEY IDEA", "Every instruction word (FROM, RUN, COPY, WORKDIR, USER, "
        "ENTRYPOINT...) is a directive TO Docker. 'RUN something' means \u201cDocker, run "
        "'something' inside the container while building.\u201d It does not mean you run it.")
callout("[!] WATCH OUT", "This particular Dockerfile assumes a Go application (it expects "
        "go.mod, go.sum and Go source). It is a best-practice TEMPLATE. If you point "
        "'docker build' at it without a real Go app present, the 'COPY go.mod go.sum ./' "
        "line will fail. To try it end to end, use a Go app (or swap the builder stage "
        "for your app's language). The nginx app in examples/ is a separate, simpler demo.")

h2("3.2 Stage 1 \u2014 the builder")
body("This is the FIRST of three stages inside the single Dockerfile. Everything from "
     "this 'FROM' line down to the next 'FROM' line is Stage 1. Its job: compile your "
     "code and run your tests inside a full Go toolchain. None of it survives into the "
     "final image \u2014 only the compiled binary is carried forward.")
code_block(
    "FROM --platform=$BUILDPLATFORM golang:1.21-alpine AS builder\n\n"
    "RUN apk add --no-cache git ca-certificates\n"
    "WORKDIR /app\n\n"
    "COPY go.mod go.sum ./\n"
    "RUN go mod download && go mod verify\n\n"
    "COPY . .\n\n"
    "RUN CGO_ENABLED=0 GOOS=linux \\\n"
    "    go build -ldflags=\"-w -s -trimpath\" -o /server .\n\n"
    "RUN CGO_ENABLED=0 GOOS=linux go test -v -cover ./...")
bullet("**FROM ... AS builder** \u2014 names this stage so later stages can copy from it. "
       "**$BUILDPLATFORM** lets Buildx build natively for the host arch then cross-compile.")
bullet("**apk add --no-cache** \u2014 Alpine's package manager; --no-cache avoids leaving an "
       "index behind (smaller layer, nothing to clean up).")
bullet("**COPY go.mod go.sum first, THEN run download** \u2014 this is deliberate layer "
       "caching. Dependencies change rarely; source changes often. By copying just the "
       "module files first, Docker reuses the cached dependency layer until go.mod itself "
       "changes \u2014 builds become dramatically faster.")
bullet("**CGO_ENABLED=0** \u2014 disables C bindings so Go produces a fully static binary with "
       "no shared-library dependencies. This is what lets us later run on 'scratch' (an "
       "empty image with no libc at all).")
bullet("**-ldflags=\"-w -s -trimpath\"** \u2014 -w/-s strip debug symbols (smaller binary, less "
       "info for an attacker); -trimpath removes local filesystem paths from the binary.")
bullet("**Tests run inside the build** \u2014 if a test fails, the image never gets built. The "
       "build itself is your first quality gate.")
callout("[>] WHAT YOU ACTUALLY DO HERE", "Nothing line-by-line. You do NOT type 'FROM', "
        "'RUN go build', etc. These instructions live in dockerfiles/Dockerfile and Docker "
        "executes them for you inside the build. Your only action for this entire stage is "
        "the single 'docker build' command (see Part 4) run from the repo root in Git-Bash. "
        "You touch THIS section only to change how your app compiles \u2014 e.g. swap the Go "
        "version on the 'FROM' line, or add a build flag \u2014 then save and re-run 'docker build'.")

h2("3.3 Stage 2 \u2014 certificates")
body("Everything from this second 'FROM' line down to the next 'FROM' is Stage 2. It is "
     "a tiny, separate stage whose only job is to grab a clean set of TLS root "
     "certificates that the final image will need for HTTPS.")
code_block(
    "FROM alpine:3.19 AS certs\n"
    "RUN apk --update add ca-certificates")
body("A separate tiny stage whose only job is to provide the TLS root certificates your "
     "app needs to make HTTPS calls. Isolating it keeps the final image minimal and the "
     "intent clear.")
bullet("**FROM alpine:3.19 AS certs** \u2014 starts a fresh stage named 'certs' from a "
       "minimal Alpine Linux. Naming it lets Stage 3 copy from it by name.")
bullet("**apk --update add ca-certificates** \u2014 installs the standard bundle of trusted "
       "certificate authorities into /etc/ssl/certs/ inside this stage.")
callout("[>] WHAT YOU ACTUALLY DO HERE", "Nothing by hand. You do not type these two "
        "lines. They already live in dockerfiles/Dockerfile, and Docker runs them "
        "automatically as part of the same 'docker build' command from Stage 1. This "
        "stage produces no output you interact with \u2014 it just sits ready for Stage 3 to "
        "copy the certs out of it. You only ever touch this section if you want to CHANGE "
        "it (e.g. bump the Alpine version); then you edit the file and re-run 'docker build'.")

h2("3.4 Stage 3 \u2014 the production image")
body("The third and final 'FROM' begins Stage 3 \u2014 the image you actually ship. It starts "
     "from nothing and pulls in only the two finished artifacts from the earlier stages: "
     "the compiled binary (from 'builder') and the certificates (from 'certs').")
code_block(
    "FROM scratch\n\n"
    "COPY --from=certs /etc/ssl/certs/ca-certificates.crt /etc/ssl/certs/\n"
    "COPY --from=builder /server /server\n\n"
    "USER 65534:65534\n\n"
    "HEALTHCHECK --interval=30s --timeout=3s --start-period=5s --retries=3 \\\n"
    "    CMD [\"/server\", \"/health\"]\n\n"
    "EXPOSE 8080\n"
    "ENTRYPOINT [\"/server\"]")
bullet("**FROM scratch** \u2014 the empty base image. Zero OS, zero shell, zero package "
       "manager. Smallest possible attack surface. Only your static binary + certs exist.")
bullet("**COPY --from=...** \u2014 pulls exactly two artifacts forward from the earlier stages. "
       "Nothing else from those stages survives.")
bullet("**USER 65534:65534** \u2014 the 'nobody' user. The container runs unprivileged, so even "
       "if compromised, the process has no root rights.")
bullet("**HEALTHCHECK** \u2014 Docker/orchestrators periodically run this; a failing health "
       "check marks the container unhealthy and can trigger a restart.")
bullet("**ENTRYPOINT (exec form)** \u2014 the JSON-array form runs the binary directly as PID 1 "
       "without a shell wrapper, so signals (SIGTERM on shutdown) reach your app cleanly.")
callout("[!] PITFALL", "Because 'scratch' has no shell, you cannot 'docker exec' a bash "
        "session into this container to debug it. For debugging, build the 'builder' "
        "target instead (see the dev compose file in Part 4).")
callout("[>] WHAT YOU ACTUALLY DO HERE", "Again, nothing by hand \u2014 Docker assembles this "
        "final image automatically at the end of the same 'docker build' run. The PAYOFF "
        "of this stage is the image itself: after the build finishes you interact with it "
        "with normal commands like 'docker run -p 8080:8080 myapp:latest' (run it) or "
        "'docker images' (see it). You edit THIS section only to change runtime behaviour "
        "\u2014 the exposed port, the health check, or the user the app runs as.")

# ============================================================
# PART 4 - BUILD LOCALLY
# ============================================================
h1("Part 4 \u2014 Building & Testing Locally")
h2("4.1 A simple build")
code_block(
    "# from the repo root\n"
    "docker build -t myapp:latest -f dockerfiles/Dockerfile .")
bullet("**-t myapp:latest** \u2014 tags (names) the resulting image.")
bullet("**-f dockerfiles/Dockerfile** \u2014 path to the Dockerfile.")
bullet("**.** (the final dot) \u2014 the build context: the directory Docker sends to the "
       "builder. .dockerignore controls what gets excluded.")

h2("4.2 Multi-architecture builds with Buildx")
body("Real platforms serve both Intel/AMD (amd64) and ARM (arm64, e.g. Apple Silicon, "
     "AWS Graviton). Buildx + QEMU emulation builds both from one machine:")
code_block(
    "docker buildx create --use --name multiarch\n"
    "docker buildx build \\\n"
    "  --platform linux/amd64,linux/arm64 \\\n"
    "  -t myapp:latest \\\n"
    "  -f dockerfiles/Dockerfile \\\n"
    "  --load .")
callout("[i] NOTE", "--load only works for a single platform locally; to keep BOTH "
        "architectures you --push to a registry, which stores them under one tag as a "
        "'manifest list'. The CI pipeline does exactly this.")

h2("4.3 The development compose file")
body("dockerfiles/docker-compose.dev.yml targets the 'builder' stage (which has a shell "
     "and tools) and mounts your source as a volume for fast local iteration:")
code_block(
    "services:\n"
    "  app:\n"
    "    build:\n"
    "      context: ..\n"
    "      dockerfile: dockerfiles/Dockerfile\n"
    "      target: builder        # <- the debuggable stage\n"
    "    ports: [\"8080:8080\"]\n"
    "    environment:\n"
    "      - ENVIRONMENT=development\n"
    "      - LOG_LEVEL=debug\n"
    "    volumes:\n"
    "      - ../:/app")
code_block("docker compose -f dockerfiles/docker-compose.dev.yml up --build")

# ============================================================
# PART 5 - TAGGING
# ============================================================
h1("Part 5 \u2014 Image Tagging Strategy")
body("A tag is a human-friendly pointer to an image digest (a sha256 hash). The same "
     "image gets MULTIPLE tags so different consumers can refer to it the way that suits "
     "them.")
add_image("diagram_tagging.png", width=6.4,
          caption="Figure 3 \u2014 One immutable digest, many tags pointing at it.")
info_table(
    ["Tag type", "Example", "Use case"],
    [
        ["latest", "myapp:latest", "Dev \u2014 always newest on main"],
        ["semver", "myapp:1.4.2", "Production releases (immutable)"],
        ["range", "myapp:1.4 / myapp:1", "Staging \u2014 newest patch in a line"],
        ["commit SHA", "myapp:abc1234", "Reproducible / traceable builds"],
        ["timestamp", "myapp:20260630-abc1234", "Audit trail"],
    ])

h2("5.1 parse-version.sh \u2014 extracting a semantic version")
code_block(
    "REF=\"${1:-$(git describe --tags --always --abbrev=7 ...)}\"\n"
    "REF=\"${REF#v}\"                       # strip leading 'v'\n"
    "if [[ \"$REF\" =~ ^([0-9]+)\\.([0-9]+)\\.([0-9]+)$ ]]; then\n"
    "    echo \"$REF\"; exit 0              # it's a clean semver tag\n"
    "fi\n"
    "echo \"0.0.0-dev-$(git rev-parse --short HEAD)\"  # fallback")
body("This reads the latest git tag. If it looks like 1.4.2 it returns it; otherwise it "
     "produces a safe dev version stamped with the commit SHA. The regex with "
     "BASH_REMATCH is how bash captures the major/minor/patch groups.")

h2("5.2 image-tagging.sh \u2014 generating all the tags")
code_block(
    "SHA=$(git rev-parse --short HEAD)\n"
    "DATE=$(date +%Y%m%d)\n"
    "VERSION=$(git describe --tags --abbrev=0 | sed 's/^v//' || echo 0.0.0)\n\n"
    "case \"$TAG_TYPE\" in\n"
    "  semver) tags+=(\"$IMAGE:$VERSION\")\n"
    "          tags+=(\"$IMAGE:${VERSION%.*}\")   # 1.4  (strip .patch)\n"
    "          tags+=(\"$IMAGE:${VERSION%%.*}\")  # 1    (strip .minor.patch)\n"
    "esac")
bullet("**${VERSION%.*}** \u2014 bash parameter expansion: remove the shortest trailing '.*', "
       "turning 1.4.2 into 1.4.")
bullet("**${VERSION%%.*}** \u2014 remove the longest trailing '.*', turning 1.4.2 into 1.")
callout("[!] BUG SPOTTED", "In your image-tagging.sh the 'case' has 'all|latest' first AND "
        "'semver|all' second. Because bash 'case' stops at the first match, the 'all' "
        "option only ever hits the latest branch \u2014 semver/sha/timestamp tags are skipped "
        "when TAG_TYPE=all. Fix: give 'all' its own branch that appends every tag type, or "
        "remove 'all' from the latest line. (See Appendix A.)")

# ============================================================
# PART 6 - SCANNING
# ============================================================
h1("Part 6 \u2014 Vulnerability Scanning with Trivy")
body("Trivy inspects your image's OS packages and language dependencies against CVE "
     "databases, and also checks for misconfigurations. It is the security gate of the "
     "pipeline.")
h2("6.1 Local scan")
code_block(
    "trivy image myapp:latest\n\n"
    "# only fail on the serious stuff:\n"
    "trivy image --severity HIGH,CRITICAL --ignore-unfixed myapp:latest")
bullet("**--ignore-unfixed** \u2014 skip CVEs that have no patch available yet (you can't fix "
       "them, so don't fail the build on them).")
bullet("**--severity** \u2014 which levels to report.")
h2("6.2 The gating philosophy")
body("Your README defines the policy: WARN on High, FAIL on Critical. This is the "
     "pragmatic middle ground \u2014 criticals block the release; highs are surfaced but don't "
     "halt everything. In ci.yml this maps to two Trivy steps: one uploads a full SARIF "
     "report (exit-code 0, never blocks) for the GitHub Security tab, and a second config "
     "scan with exit-code 1 that actually fails the job.")
code_block(
    "- name: Run Trivy vulnerability scanner\n"
    "  uses: aquasecurity/trivy-action@master\n"
    "  with:\n"
    "    image-ref: '${{ env.REGISTRY }}/${{ env.IMAGE_NAME }}:${{ github.sha }}'\n"
    "    format: 'sarif'\n"
    "    output: 'trivy-results.sarif'\n"
    "    severity: 'CRITICAL,HIGH,MEDIUM'\n"
    "    exit-code: '0'          # report only, never blocks")

# ============================================================
# PART 7 - CI PIPELINE
# ============================================================
h1("Part 7 \u2014 The GitHub Actions CI Pipeline")
body("Open .github/workflows/ci.yml. This is the automation that ties everything "
     "together. We'll go job by job.")
add_image("diagram_jobgraph.png", width=6.4,
          caption="Figure 4 \u2014 Jobs chained with 'needs:'. A failure stops everything downstream.")

h2("7.1 Triggers, env, and permissions")
code_block(
    "on:\n"
    "  push:\n"
    "    branches: [main, master]\n"
    "    tags: ['v*']\n"
    "  pull_request:\n"
    "    branches: [main, master]\n\n"
    "env:\n"
    "  REGISTRY: ghcr.io\n"
    "  IMAGE_NAME: ${{ github.repository }}\n"
    "  BUILD_PLATFORMS: linux/amd64,linux/arm64\n\n"
    "permissions:\n"
    "  contents: read\n"
    "  packages: write\n"
    "  security-events: write")
bullet("**on.push.tags ['v*']** \u2014 pushing a tag like v1.4.2 triggers a release build.")
bullet("**permissions** \u2014 least privilege for the workflow token: read code, write "
       "packages, upload security results. Nothing more.")

h2("7.2 Job 1 \u2014 build (and test)")
body("Sets up QEMU + Buildx (multi-arch), computes all tags with docker/metadata-action, "
     "logs in, then builds and pushes:")
code_block(
    "- uses: docker/metadata-action@v5\n"
    "  with:\n"
    "    tags: |\n"
    "      type=sha,prefix=,format=long\n"
    "      type=semver,pattern={{version}}\n"
    "      type=semver,pattern={{major}}.{{minor}}\n"
    "      type=raw,value=latest,enable={{is_default_branch}}\n\n"
    "- uses: docker/build-push-action@v5\n"
    "  with:\n"
    "    platforms: ${{ env.BUILD_PLATFORMS }}\n"
    "    push: ${{ github.event_name != 'pull_request' }}\n"
    "    cache-from: type=gha\n"
    "    cache-to: type=gha,mode=max\n"
    "    provenance: mode=max\n"
    "    sbom: true")
bullet("**push only when NOT a pull_request** \u2014 PRs build to validate but never publish. "
       "This prevents untrusted PR code from pushing images.")
bullet("**cache-from/to: type=gha** \u2014 uses GitHub's cache so layers persist between runs.")
bullet("**provenance + sbom** \u2014 generates supply-chain attestations automatically.")

h2("7.3 Job 2 \u2014 scan")
body("needs: [build]. Runs the two Trivy steps from Part 6 and uploads SARIF to the "
     "Security tab.")
h2("7.4 Job 3 \u2014 push-ecr")
body("needs: [build, scan] and only runs on non-PR events. Configures AWS credentials, "
     "logs in to ECR, retags the image for ECR, pushes, generates an SBOM, and uploads "
     "it to S3. Uses the 'production' GitHub Environment so you can require approvals.")
h2("7.5 Job 4 \u2014 sign")
body("needs: [push-ecr]. Installs Cosign and signs the pushed image (Part 10).")
h2("7.6 Job 5 \u2014 notify")
body("if: always(). Writes a markdown summary table of every job's result to the run "
     "summary so you get a one-glance status report.")
callout("[!] BUG SPOTTED", "In the push-ecr 'docker tag' lines, the latest tag uses "
        "$ECR_REPO/$ECR_REPO:latest instead of $ECR_REGISTRY/$ECR_REPO:latest. The push "
        "target is correct but the tag source is malformed \u2014 the latest push will fail. "
        "Fix in Appendix A.")

# ============================================================
# PART 8 - AWS SETUP
# ============================================================
h1("Part 8 \u2014 AWS Setup: IAM, ECR, S3 & Secrets")
add_image("diagram_aws.png", width=6.4,
          caption="Figure 5 \u2014 How the pipeline authenticates to and uses AWS services.")
h2("8.1 Create the ECR repository")
code_block(
    "aws ecr create-repository \\\n"
    "  --repository-name myapp \\\n"
    "  --image-scanning-configuration scanOnPush=true \\\n"
    "  --region us-east-1")
bullet("**scanOnPush=true** \u2014 ECR runs its own basic scan on every push (defense in depth "
       "alongside Trivy).")

h2("8.2 The IAM permissions the pipeline needs")
body("Create an IAM role/user with a least-privilege policy. The pipeline must be able "
     "to authenticate to ECR, push images, write SBOMs to S3, and manage lifecycle:")
code_block(
    "{\n"
    "  \"Version\": \"2012-10-17\",\n"
    "  \"Statement\": [\n"
    "    {\"Effect\":\"Allow\",\n"
    "     \"Action\":[\"ecr:GetAuthorizationToken\"],\n"
    "     \"Resource\":\"*\"},\n"
    "    {\"Effect\":\"Allow\",\n"
    "     \"Action\":[\"ecr:BatchCheckLayerAvailability\",\n"
    "                \"ecr:PutImage\",\"ecr:InitiateLayerUpload\",\n"
    "                \"ecr:UploadLayerPart\",\"ecr:CompleteLayerUpload\",\n"
    "                \"ecr:BatchDeleteImage\",\"ecr:PutLifecyclePolicy\"],\n"
    "     \"Resource\":\"arn:aws:ecr:us-east-1:<ACCOUNT_ID>:repository/myapp\"},\n"
    "    {\"Effect\":\"Allow\",\n"
    "     \"Action\":[\"s3:PutObject\"],\n"
    "     \"Resource\":\"arn:aws:s3:::<SBOM_BUCKET>/*\"}\n"
    "  ]\n"
    "}")
callout("[+] BEST PRACTICE", "Prefer GitHub OIDC over long-lived AWS keys: GitHub mints a "
        "short-lived token per run and assumes the role, so there are no static secrets to "
        "leak. Your sister 3-tier project already uses OIDC \u2014 reuse that pattern here.")

h2("8.3 The GitHub Secrets you must set")
body("Repository \u2192 Settings \u2192 Secrets and variables \u2192 Actions. ci.yml references all of "
     "these:")
info_table(
    ["Secret", "Purpose"],
    [
        ["AWS_ACCESS_KEY_ID", "AWS auth (or use OIDC role instead)"],
        ["AWS_SECRET_ACCESS_KEY", "AWS auth secret"],
        ["AWS_REGION", "e.g. us-east-1"],
        ["ECR_REGISTRY", "<account>.dkr.ecr.<region>.amazonaws.com"],
        ["ECR_REPO", "Repository name, e.g. myapp"],
        ["SBOM_BUCKET", "S3 bucket name for SBOM storage"],
        ["COSIGN_YAML", "Cosign signing config/key material"],
    ])
callout(" SECURITY", "Never hardcode these in the YAML. They are injected as masked "
        "environment values at run time \u2014 this is the 'security tool / secrets in an env "
        "var' requirement from your project brief.")

# ============================================================
# PART 9 - ECR PUSH + CODEBUILD
# ============================================================
h1("Part 9 \u2014 Pushing to ECR & the CodeBuild Alternative")
h2("9.1 Manual push (understand what CI automates)")
code_block(
    "aws ecr get-login-password --region us-east-1 | \\\n"
    "  docker login --username AWS --password-stdin \\\n"
    "  <account>.dkr.ecr.us-east-1.amazonaws.com\n\n"
    "docker tag myapp:latest \\\n"
    "  <account>.dkr.ecr.us-east-1.amazonaws.com/myapp:1.4.2\n"
    "docker push <account>.dkr.ecr.us-east-1.amazonaws.com/myapp:1.4.2")
body("get-login-password returns a temporary token; piping it into docker login "
     "authenticates without storing credentials on disk.")

h2("9.2 The CodeBuild alternative (.github/codebuild-spec.yml)")
body("Your project brief asks to 'build using EC2 and put it in the container'. AWS "
     "CodeBuild runs your build on AWS-managed EC2 compute. The buildspec mirrors the CI "
     "pipeline in three phases:")
code_block(
    "version: 0.2\n"
    "phases:\n"
    "  pre_build:\n"
    "    commands:\n"
    "      - aws ecr get-login-password --region $AWS_REGION | \\\n"
    "          docker login --username AWS --password-stdin $ECR_REGISTRY\n"
    "      - COMMIT_HASH=$(echo $CODEBUILD_RESOLVED_SOURCE_VERSION | cut -c 1-7)\n"
    "      - IMAGE_TAG=${COMMIT_HASH}\n"
    "  build:\n"
    "    commands:\n"
    "      - docker build --platform linux/amd64,linux/arm64 \\\n"
    "          --cache-from $ECR_REGISTRY/$ECR_REPO:latest \\\n"
    "          -t $ECR_REGISTRY/$ECR_REPO:$IMAGE_TAG -f dockerfiles/Dockerfile .\n"
    "  post_build:\n"
    "    commands:\n"
    "      - trivy image --severity HIGH,CRITICAL $ECR_REGISTRY/$ECR_REPO:$IMAGE_TAG\n"
    "      - docker push $ECR_REGISTRY/$ECR_REPO:$IMAGE_TAG")
bullet("**pre_build** \u2014 authenticate + compute the tag from the commit hash CodeBuild "
       "exposes in $CODEBUILD_RESOLVED_SOURCE_VERSION.")
bullet("**build** \u2014 the docker build, reusing :latest as a cache source for speed.")
bullet("**post_build** \u2014 scan, then push. Artifacts (image-listing.json) are exported.")
callout("[i] NOTE", "GitHub Actions and CodeBuild are two front-ends to the SAME Dockerfile "
        "and scanning logic. Pick one as primary; the other demonstrates portability.")

# ============================================================
# PART 10 - SIGNING
# ============================================================
h1("Part 10 \u2014 Image Signing & SBOMs with Cosign")
h2("10.1 Why sign?")
body("A signature cryptographically proves an image came from your pipeline and hasn't "
     "been tampered with since. Combined with an SBOM (the ingredient list), you get "
     "supply-chain integrity \u2014 a core DevSecOps requirement.")
h2("10.2 signing.yml \u2014 triggered after a successful build")
code_block(
    "on:\n"
    "  workflow_run:\n"
    "    workflows: [\"Container Build Pipeline\"]\n"
    "    types: [completed]\n"
    "jobs:\n"
    "  sign-image:\n"
    "    if: ${{ github.event.workflow_run.conclusion == 'success' }}")
body("It looks up the freshly pushed image's digest (signing by digest, not tag, is the "
     "secure way \u2014 tags can move, digests cannot), then:")
code_block(
    "IMAGE=\"$ECR_REGISTRY/$ECR_REPO@$DIGEST\"\n"
    "cosign sign --yes \"$IMAGE\"\n\n"
    "trivy image --format spdx-json --output sbom.spdx.json \"$IMAGE\"\n"
    "cosign attach sbom --sbom sbom.spdx.json --type spdx \"$IMAGE\"")
h2("10.3 Verifying a signature")
code_block("cosign verify --key cosign.pub \\\n"
           "  <account>.dkr.ecr.us-east-1.amazonaws.com/myapp@sha256:...")
callout("[!] PITFALL", "Signing by tag is unsafe \u2014 always sign and verify by @sha256 "
        "digest, which is exactly what signing.yml does.")

# ============================================================
# PART 11 - LIFECYCLE
# ============================================================
h1("Part 11 \u2014 Image Lifecycle Management & Cleanup")
body("Registries fill up fast and storage costs money. Two mechanisms keep ECR tidy: a "
     "declarative lifecycle policy (ECR enforces it) and an imperative cleanup script "
     "(for finer control).")
h2("11.1 The lifecycle policy (registry/lifecycle-policy.json)")
info_table(
    ["Priority", "Rule"],
    [
        ["1", "Keep only the most recent 30 images; expire the rest"],
        ["2", "Expire 'latest'-prefixed images older than 90 days"],
        ["3", "Delete untagged images older than 7 days"],
    ])
body("ecr.yml applies it on a daily schedule (cron '0 2 * * *') with aws ecr "
     "put-lifecycle-policy.")

h2("11.2 cleanup-ecr.sh \u2014 the imperative cleaner")
code_block(
    "images=$(aws ecr list-images --repository-name \"$ECR_REPO\" \\\n"
    "  --query 'images[*] | sort_by(@, &imagePushTime)' --output json)\n\n"
    "# Strategy 1: remove untagged images older than 7 days\n"
    "# Strategy 2: if total > MAX_IMAGES, remove the oldest overflow\n\n"
    "aws ecr batch-delete-image --repository-name \"$ECR_REPO\" \\\n"
    "  --image-ids imageDigest=\"$digest\"")
bullet("**sort_by(@, &imagePushTime)** \u2014 a JMESPath query that sorts images oldest-first so "
       "the script can target the oldest for deletion.")
bullet("**DRY_RUN flag** \u2014 the 4th argument; when 'true' the script reports what it WOULD "
       "delete without deleting. ecr.yml defaults dry_run to 'true' \u2014 a safety net.")
callout("[+] BEST PRACTICE", "Always run cleanup with DRY_RUN=true first and read the output "
        "before doing a real deletion. Deletions are irreversible.")

# ============================================================
# PART 12 - TERRAFORM
# ============================================================
h1("Part 12 \u2014 Infrastructure as Code (Terraform)")
body("Rather than clicking in the AWS console, define the ECR repo and its lifecycle "
     "policy as code in terraform/modules/ecs-builder/. This satisfies the 'IaC' "
     "requirement and makes the infra reproducible and reviewable.")
h2("12.1 A minimal ECR module")
code_block(
    "# main.tf\n"
    "resource \"aws_ecr_repository\" \"this\" {\n"
    "  name                 = var.repo_name\n"
    "  image_tag_mutability = \"IMMUTABLE\"\n"
    "  image_scanning_configuration { scan_on_push = true }\n"
    "}\n\n"
    "resource \"aws_ecr_lifecycle_policy\" \"this\" {\n"
    "  repository = aws_ecr_repository.this.name\n"
    "  policy     = file(\"${path.module}/../../../registry/lifecycle-policy.json\")\n"
    "}")
code_block(
    "# variables.tf\n"
    "variable \"repo_name\" { type = string }\n\n"
    "# outputs.tf\n"
    "output \"repository_url\" { value = aws_ecr_repository.this.repository_url }")
bullet("**image_tag_mutability = IMMUTABLE** \u2014 once a tag points at a digest it can never "
       "be reassigned. This prevents 'tag hijacking' where 1.4.2 silently changes.")
bullet("**Reusing lifecycle-policy.json** \u2014 the same policy file feeds both the runtime "
       "ecr.yml workflow and Terraform, so there is one source of truth.")
h2("12.2 Applying it")
code_block(
    "cd terraform\n"
    "terraform init      # download providers, configure backend\n"
    "terraform plan      # preview changes (read-only)\n"
    "terraform apply     # create the ECR repo + policy")

# ============================================================
# PART 13 - END TO END
# ============================================================
h1("Part 13 \u2014 End-to-End Run-Through & Verification")
body("Put it all together. This is the full loop from code change to signed image:")
numbered("Make a code change and commit it on a feature branch.")
numbered("Open a pull request \u2192 ci.yml builds + scans but does NOT push (PR safety).")
numbered("Merge to main \u2192 ci.yml builds multi-arch, scans, pushes to ECR, writes SBOM to "
         "S3, then signing.yml signs the image by digest.")
numbered("Tag a release: git tag v1.4.2 && git push --tags \u2192 semver tags are published.")
numbered("Nightly \u2192 ecr.yml applies the lifecycle policy and runs cleanup (dry-run).")
h2("13.1 Verification checklist")
code_block(
    "# image is in ECR with expected tags\n"
    "aws ecr describe-images --repository-name myapp --region us-east-1\n\n"
    "# SBOM landed in S3\n"
    "aws s3 ls s3://<SBOM_BUCKET>/\n\n"
    "# signature verifies\n"
    "cosign verify --key cosign.pub <account>.dkr.ecr.us-east-1.amazonaws.com/myapp@sha256:...\n\n"
    "# pull and run\n"
    "docker run --rm -p 8080:8080 <account>.dkr.ecr.us-east-1.amazonaws.com/myapp:1.4.2")

# ============================================================
# APPENDIX A - TROUBLESHOOTING
# ============================================================
h1("Appendix A \u2014 Troubleshooting & Known Issues in This Repo")
h2("A.1 Two real bugs found in your current code")
body("image-tagging.sh \u2014 the 'all' case never emits semver/sha/timestamp tags. Fix the "
     "case block:", bold=True)
code_block(
    "case \"$TAG_TYPE\" in\n"
    "  all)\n"
    "    tags+=(\"$IMAGE_NAME:latest\")\n"
    "    tags+=(\"$IMAGE_NAME:$VERSION\")\n"
    "    tags+=(\"$IMAGE_NAME:${VERSION%.*}\")\n"
    "    tags+=(\"$IMAGE_NAME:${VERSION%%.*}\")\n"
    "    tags+=(\"$IMAGE_NAME:$SHA\")\n"
    "    tags+=(\"$IMAGE_NAME:${DATE}-${SHA}\")\n"
    "    ;;\n"
    "  latest)    tags+=(\"$IMAGE_NAME:latest\") ;;\n"
    "  semver)    tags+=(\"$IMAGE_NAME:$VERSION\" \"$IMAGE_NAME:${VERSION%.*}\" \"$IMAGE_NAME:${VERSION%%.*}\") ;;\n"
    "  sha)       tags+=(\"$IMAGE_NAME:$SHA\") ;;\n"
    "  timestamp) tags+=(\"$IMAGE_NAME:${DATE}-${SHA}\") ;;\n"
    "esac")
body("ci.yml push-ecr \u2014 the 'latest' docker tag source is wrong:", bold=True)
code_block(
    "# WRONG (current):\n"
    "docker tag ...:$IMAGE_TAG $ECR_REPO/$ECR_REPO:latest\n\n"
    "# CORRECT:\n"
    "docker tag $ECR_REGISTRY/$ECR_REPO:$IMAGE_TAG $ECR_REGISTRY/$ECR_REPO:latest")

h2("A.2 Common errors")
info_table(
    ["Symptom", "Likely cause & fix"],
    [
        ["denied: not authorized (ECR)", "Token expired or wrong region. Re-run get-login-password."],
        ["exec /server: no such file", "Binary built for wrong arch, or not static. Check CGO_ENABLED=0."],
        ["Trivy job fails the build", "A CRITICAL CVE exists. Patch the base image or add --ignore-unfixed."],
        ["multi-arch build fails locally", "QEMU not set up. Run docker/setup-qemu-action or install qemu."],
        ["cosign verify fails", "Verifying by tag not digest, or wrong public key."],
        [".sh scripts 'not found' on Windows", "Run in Git-Bash, and chmod +x the scripts."],
    ])

# ============================================================
# APPENDIX B - GLOSSARY
# ============================================================
h1("Appendix B \u2014 Glossary")
info_table(
    ["Term", "Definition"],
    [
        ["BuildKit", "Docker's modern build engine; enables caching, secrets, multi-arch."],
        ["Buildx", "CLI plugin front-end to BuildKit for multi-platform builds."],
        ["Manifest list", "An index pointing at per-architecture images under one tag."],
        ["Digest", "Immutable sha256 content hash uniquely identifying an image."],
        ["OIDC", "OpenID Connect; lets GitHub assume an AWS role without static keys."],
        ["SARIF", "Standard format for static-analysis results (the GitHub Security tab)."],
        ["JMESPath", "Query language for JSON used in AWS CLI --query expressions."],
        ["Distroless", "Minimal base images with no shell/package manager (alt. to scratch)."],
        ["PID 1", "The first process in a container; must handle signals for clean shutdown."],
    ])

doc.add_paragraph()
end = doc.add_paragraph()
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = end.add_run("\u2014 End of Tutorial \u2014\nContainer Build Platform \u00b7 Enterprise Container Factory")
r.italic = True; r.font.color.rgb = GRAY; r.font.size = Pt(10)

doc.core_properties.author = "Blakeman, Daniel"
doc.core_properties.last_modified_by = "Blakeman, Daniel"
doc.save(OUT)
print("Saved:", OUT, os.path.getsize(OUT), "bytes")
